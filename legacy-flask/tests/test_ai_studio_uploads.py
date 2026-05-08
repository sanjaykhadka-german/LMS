"""Regression tests for AI Studio upload + chat flow.

Locks in the cookie-overflow fix: extracted .docx text now lives in a sibling
.txt file on disk, not in `session["ai_studio"]["files"][id]["extracted_text"]`.
Without that fix, a few SOPs would push Flask's signed-cookie session past the
browser's ~4 KB limit and the next /message request would see an empty
files_map, causing Claude to reply as if no files were attached.

All tests drive the real Flask routes via `auth_client` (admin) and
monkeypatch `claude_service._call_claude` so no network call is made.
"""
import io
import json
import os

import docx
import pytest


pytestmark = [pytest.mark.regression, pytest.mark.backend]


# -----------------------------------------------------------------------------
# Fixtures
# -----------------------------------------------------------------------------

@pytest.fixture(autouse=True)
def _ai_studio_env(app, tmp_path, monkeypatch):
    """Per-test: route uploads to tmp_path, force Claude as the AI provider."""
    monkeypatch.setitem(app.config, "UPLOAD_FOLDER", str(tmp_path))
    monkeypatch.setitem(app.config, "CLAUDE_API_KEY", "test-key")
    monkeypatch.setitem(app.config, "AI_PROVIDER", "claude")


def _make_docx(text):
    """Build a real .docx in-memory whose paragraph text == `text`."""
    buf = io.BytesIO()
    doc = docx.Document()
    # add_paragraph caps practical paragraph length but accepts long strings;
    # split to keep python-docx happy on huge bodies.
    for chunk_start in range(0, len(text), 5000):
        doc.add_paragraph(text[chunk_start:chunk_start + 5000])
    doc.save(buf)
    buf.seek(0)
    return buf


def _post_upload(client, filename, payload_bytes, mimetype):
    return client.post(
        "/admin/modules/ai-studio/upload",
        data={"file": (io.BytesIO(payload_bytes), filename)},
        content_type="multipart/form-data",
    )


def _upload_docx(client, filename, body_text):
    buf = _make_docx(body_text)
    return _post_upload(
        client, filename, buf.getvalue(),
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )


def _session_cookie_size(client):
    cookie = client.get_cookie("session")
    return len(cookie.value) if cookie else 0


# -----------------------------------------------------------------------------
# Tests
# -----------------------------------------------------------------------------

def test_upload_single_docx_writes_sibling_txt(auth_client, tmp_path):
    body = "Cleaning SOP. Step 1: rinse. Step 2: sanitise. " * 50
    resp = _upload_docx(auth_client, "sop.docx", body)
    assert resp.status_code == 200, resp.data
    data = resp.get_json()
    assert data["kind"] == "docx"
    file_id = data["file_id"]

    with auth_client.session_transaction() as sess:
        meta = sess["ai_studio"]["files"][file_id]
        assert "text_path" in meta
        assert "extracted_text" not in meta, (
            "extracted_text in session would overflow cookie when many docx "
            "are uploaded — must live on disk only"
        )
        text_path = meta["text_path"]
        local_path = meta["local_path"]

    assert os.path.exists(text_path), "sibling .txt should be written at upload"
    assert os.path.exists(local_path)
    on_disk = open(text_path, encoding="utf-8").read()
    assert "Step 1: rinse" in on_disk
    assert "Step 2: sanitise" in on_disk


def test_upload_many_docx_keeps_session_cookie_under_4kb(auth_client):
    """The actual regression: 6 chunky SOPs would push the cookie past 4 KB
    if extracted_text were still in the session. Each docx here has ~20 KB of
    text — total ~120 KB if it leaked into the cookie."""
    big_body = ("Allergen control procedure. " * 800)  # ~22 KB
    assert len(big_body) > 20_000

    for i in range(6):
        resp = _upload_docx(auth_client, f"sop_{i}.docx", big_body)
        assert resp.status_code == 200, resp.data
        size = _session_cookie_size(auth_client)
        assert size < 4000, (
            f"session cookie is {size} bytes after upload {i + 1} — "
            "browser will silently drop the Set-Cookie and the next /message "
            "request will see an empty files_map (the original bug)"
        )

    with auth_client.session_transaction() as sess:
        assert len(sess["ai_studio"]["files"]) == 6


def test_message_with_multiple_docx_sends_all_to_claude(auth_client, monkeypatch):
    """Pre-fix symptom: Claude was called with only the user's typed text and
    no file blocks, so it asked for a source document. Verify every uploaded
    docx becomes its own text block in the request to Claude."""
    import claude_service

    file_ids = []
    bodies = {}
    for i, name in enumerate(["cleaning", "allergen", "glass", "spillage"]):
        body = f"{name.upper()} PROCEDURE — unique marker {name}-{i}. " * 30
        bodies[name] = body
        resp = _upload_docx(auth_client, f"{name}.docx", body)
        assert resp.status_code == 200
        file_ids.append(resp.get_json()["file_id"])

    captured = {}

    def fake_call(api_key, body):
        captured["body"] = body
        return (
            "Building the module from your SOPs.\n\n"
            '```json\n{"title": "Combined SOP", "sections": [], "quiz": '
            '{"questions": []}}\n```'
        )

    monkeypatch.setattr(claude_service, "_call_claude", fake_call)

    resp = auth_client.post(
        "/admin/modules/ai-studio/message",
        data=json.dumps({"message": "build me a module", "file_ids": file_ids}),
        content_type="application/json",
    )
    assert resp.status_code == 200, resp.data
    assert resp.get_json()["module_json"], "fake reply contained a JSON block"

    sent_messages = captured["body"]["messages"]
    assert len(sent_messages) == 1
    blocks = sent_messages[0]["content"]
    text_blocks = [b for b in blocks if b.get("type") == "text"]

    # 4 docx attachments + 1 user-typed message
    assert len(text_blocks) == 5, (
        f"expected 4 file blocks + 1 user-text block, got {len(text_blocks)}"
    )
    joined = "\n".join(b["text"] for b in text_blocks)
    for name, body in bodies.items():
        assert f"[Attached .docx: {name}.docx]" in joined
        # the unique marker must reach Claude — proves disk read worked
        assert f"unique marker {name}" in joined
    assert "build me a module" in joined


def test_upload_doc_returns_clean_400(auth_client, caplog):
    """Legacy .doc rejection is a clean ValueError → 400, no traceback."""
    resp = _post_upload(
        auth_client, "old.doc", b"\xd0\xcf\x11\xe0fake-ole-header",
        "application/msword",
    )
    assert resp.status_code == 400
    msg = resp.get_json()["error"]
    assert ".doc" in msg
    assert "PDF" in msg or "DOCX" in msg
    # ValueError handled cleanly — no exception logged
    assert not any("Traceback" in r.message for r in caplog.records)


def test_upload_mixed_pdf_and_docx(auth_client):
    """Mixing PDFs + DOCX must keep cookie under 4 KB and stage all four with
    correct kinds. PDFs must NOT be base64-stuffed into the session."""
    minimal_pdf = (
        b"%PDF-1.4\n"
        b"1 0 obj<</Type/Catalog/Pages 2 0 R>>endobj\n"
        b"2 0 obj<</Type/Pages/Kids[3 0 R]/Count 1>>endobj\n"
        b"3 0 obj<</Type/Page/Parent 2 0 R/MediaBox[0 0 100 100]>>endobj\n"
        b"xref\n0 4\n0000000000 65535 f\n"
        b"trailer<</Size 4/Root 1 0 R>>\nstartxref\n0\n%%EOF\n"
    )

    r1 = _upload_docx(auth_client, "doc_one.docx", "Procedure A. " * 50)
    r2 = _post_upload(auth_client, "scan_one.pdf", minimal_pdf, "application/pdf")
    r3 = _upload_docx(auth_client, "doc_two.docx", "Procedure B. " * 50)
    r4 = _post_upload(auth_client, "scan_two.pdf", minimal_pdf, "application/pdf")

    for r in (r1, r2, r3, r4):
        assert r.status_code == 200, r.data
    kinds = [r.get_json()["kind"] for r in (r1, r2, r3, r4)]
    assert kinds == ["docx", "pdf", "docx", "pdf"]

    assert _session_cookie_size(auth_client) < 4000

    with auth_client.session_transaction() as sess:
        files = sess["ai_studio"]["files"]
        assert len(files) == 4
        for meta in files.values():
            # PDFs/images must never carry bytes in the session
            assert "data" not in meta
            assert "extracted_text" not in meta
            if meta["kind"] == "pdf":
                assert "text_path" not in meta


def test_reset_cleans_up_docx_and_sibling_txt(auth_client):
    resp = _upload_docx(auth_client, "to_clear.docx", "Some content. " * 30)
    assert resp.status_code == 200
    file_id = resp.get_json()["file_id"]

    with auth_client.session_transaction() as sess:
        meta = sess["ai_studio"]["files"][file_id]
        local_path = meta["local_path"]
        text_path = meta["text_path"]
    assert os.path.exists(local_path)
    assert os.path.exists(text_path)

    resp = auth_client.post("/admin/modules/ai-studio/reset")
    assert resp.status_code == 200

    assert not os.path.exists(local_path), "docx should be deleted on /reset"
    assert not os.path.exists(text_path), "sibling .txt should be deleted on /reset"
    with auth_client.session_transaction() as sess:
        assert "ai_studio" not in sess
