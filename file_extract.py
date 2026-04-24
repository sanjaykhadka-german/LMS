"""Stage uploaded files for whichever AI provider is active.

- .docx: extracted to text locally (no provider accepts docx natively).
- PDF / image / audio / video: stored on disk. For Gemini, files >20 MB are
  uploaded to the Files API and referenced via file_data; smaller go inline.
  For Claude, everything is passed inline by the service module when needed.
- Each provider module exposes `can_handle(meta)` which is consulted here so
  unsupported kinds (e.g. audio with Claude) fail at upload time.
"""
import logging
import os
import time
import uuid

import docx
import requests
from flask import current_app

log = logging.getLogger(__name__)

INLINE_LIMIT = 20 * 1024 * 1024      # Gemini inline_data cap
HARD_LIMIT = 500 * 1024 * 1024       # reject anything bigger than this
MIN_DOCX_CHARS = 50
DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

# Extension → (mime, kind). kind ∈ {"docx","pdf","image","audio","video"}
SUPPORTED = {
    ".pdf":  ("application/pdf", "pdf"),
    ".docx": (DOCX_MIME, "docx"),
    ".png":  ("image/png", "image"),
    ".jpg":  ("image/jpeg", "image"),
    ".jpeg": ("image/jpeg", "image"),
    ".webp": ("image/webp", "image"),
    ".heic": ("image/heic", "image"),
    ".heif": ("image/heif", "image"),
    ".mp3":  ("audio/mpeg", "audio"),
    ".wav":  ("audio/wav", "audio"),
    ".m4a":  ("audio/mp4", "audio"),
    ".ogg":  ("audio/ogg", "audio"),
    ".flac": ("audio/flac", "audio"),
    ".aac":  ("audio/aac", "audio"),
    ".mp4":  ("video/mp4", "video"),
    ".mov":  ("video/quicktime", "video"),
    ".webm": ("video/webm", "video"),
    ".mpeg": ("video/mpeg", "video"),
    ".mpg":  ("video/mpeg", "video"),
    ".avi":  ("video/x-msvideo", "video"),
}
REJECT_WITH_HINT = {".doc", ".rtf", ".odt", ".xlsx", ".xls", ".pptx", ".ppt", ".txt"}


def _ai_root():
    return os.path.join(current_app.config["UPLOAD_FOLDER"], "ai")


def prepare_file(fs, user_id, provider):
    """Validate and stage an uploaded file. Consults the active provider's
    can_handle() so unsupported kinds are rejected here. Raises ValueError
    on user-correctable problems. `provider` is 'claude' or 'gemini'."""
    filename = (getattr(fs, "filename", "") or "").strip()
    if not filename:
        raise ValueError("No file uploaded.")

    ext = os.path.splitext(filename)[1].lower()
    if ext in REJECT_WITH_HINT:
        raise ValueError(
            f"Can't read {ext} directly. Save the file as PDF or DOCX and try again."
        )
    if ext not in SUPPORTED:
        raise ValueError(
            f"Unsupported file type ({ext or 'no extension'}). "
            "Allowed: PDF, DOCX, images, audio, video."
        )

    mime, kind = SUPPORTED[ext]
    file_id = uuid.uuid4().hex

    user_dir = os.path.join(_ai_root(), str(user_id))
    os.makedirs(user_dir, exist_ok=True)
    local_path = os.path.join(user_dir, f"{file_id}{ext}")
    fs.stream.seek(0)
    data = fs.stream.read()
    size = len(data)
    if size == 0:
        raise ValueError("That file looks empty.")
    if size > HARD_LIMIT:
        raise ValueError(
            f"File is too large ({size // 1024 // 1024} MB). "
            f"Max is {HARD_LIMIT // 1024 // 1024} MB."
        )
    with open(local_path, "wb") as f:
        f.write(data)

    meta = {
        "id": file_id,
        "filename": filename,
        "mime_type": mime,
        "kind": kind,
        "size": size,
        "local_path": local_path,
        "uploaded_at": int(time.time()),
    }

    if kind == "docx":
        text = _extract_docx_text(local_path).strip()
        if len(text) < MIN_DOCX_CHARS:
            try:
                os.remove(local_path)
            except OSError:
                pass
            raise ValueError(
                "This .docx looks empty or has no readable text. "
                "Try a different file or export to PDF."
            )
        meta["extracted_text"] = text

    # Consult the active provider for kind/size policy
    ok, reason = _provider_can_handle(provider, meta)
    if not ok:
        try:
            os.remove(local_path)
        except OSError:
            pass
        raise ValueError(reason)

    # Gemini-specific: big binaries go through the Files API
    if provider == "gemini" and kind != "docx" and size > INLINE_LIMIT:
        gemini = _upload_to_files_api(local_path, mime, filename)
        meta["gemini_uri"] = gemini["uri"]
        meta["gemini_name"] = gemini["name"]
    return meta


def _provider_can_handle(provider, meta):
    # Lazy import to avoid circular imports
    if provider == "claude":
        from claude_service import can_handle
    else:
        from gemini_service import can_handle
    return can_handle(meta)


def cleanup_local_files(session_data):
    """Delete per-session files from disk. Called from /reset and the reaper."""
    if not session_data:
        return
    for meta in (session_data.get("files") or {}).values():
        p = meta.get("local_path")
        if p and os.path.exists(p):
            try:
                os.remove(p)
            except OSError as e:
                log.warning("Couldn't delete %s: %s", p, e)


def reap_old_files(user_id, max_age_seconds=7 * 24 * 3600):
    """On each /ai-studio GET, delete user's orphaned files older than max_age."""
    user_dir = os.path.join(_ai_root(), str(user_id))
    if not os.path.isdir(user_dir):
        return
    cutoff = time.time() - max_age_seconds
    for name in os.listdir(user_dir):
        path = os.path.join(user_dir, name)
        try:
            if os.path.isfile(path) and os.path.getmtime(path) < cutoff:
                os.remove(path)
        except OSError:
            pass


def _extract_docx_text(path):
    doc = docx.Document(path)
    lines = [p.text for p in doc.paragraphs if p.text and p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text and c.text.strip()]
            if cells:
                lines.append(" | ".join(cells))
    return "\n".join(lines)


def _upload_to_files_api(path, mime, display_name):
    """Upload a file to Gemini Files API via the simple multipart endpoint,
    then poll until state == ACTIVE. Returns {name, uri}."""
    api_key = current_app.config.get("GEMINI_API_KEY", "")
    if not api_key:
        raise RuntimeError("Gemini API key not configured")

    size = os.path.getsize(path)
    start_url = "https://generativelanguage.googleapis.com/upload/v1beta/files"
    headers = {
        "x-goog-api-key": api_key,
        "X-Goog-Upload-Protocol": "resumable",
        "X-Goog-Upload-Command": "start",
        "X-Goog-Upload-Header-Content-Length": str(size),
        "X-Goog-Upload-Header-Content-Type": mime,
        "Content-Type": "application/json",
    }
    start_body = {"file": {"display_name": display_name}}
    r = requests.post(start_url, headers=headers, json=start_body, timeout=60)
    if r.status_code >= 300:
        raise RuntimeError(f"Files API start failed: HTTP {r.status_code}: {r.text[:300]}")
    upload_url = r.headers.get("X-Goog-Upload-URL")
    if not upload_url:
        raise RuntimeError("Files API did not return upload URL")

    with open(path, "rb") as f:
        chunk = f.read()
    up_headers = {
        "Content-Length": str(size),
        "X-Goog-Upload-Offset": "0",
        "X-Goog-Upload-Command": "upload, finalize",
    }
    r2 = requests.post(upload_url, headers=up_headers, data=chunk, timeout=600)
    if r2.status_code >= 300:
        raise RuntimeError(f"Files API upload failed: HTTP {r2.status_code}: {r2.text[:300]}")

    body = r2.json().get("file") or {}
    file_name = body.get("name")
    file_uri = body.get("uri")
    state = body.get("state", "PROCESSING")
    if not file_name or not file_uri:
        raise RuntimeError("Files API returned no file handle")

    if state != "ACTIVE":
        _wait_for_active(file_name, api_key)
    return {"name": file_name, "uri": file_uri}


def _wait_for_active(file_name, api_key, timeout_s=120):
    url = f"https://generativelanguage.googleapis.com/v1beta/{file_name}"
    headers = {"x-goog-api-key": api_key}
    deadline = time.time() + timeout_s
    while time.time() < deadline:
        time.sleep(2)
        r = requests.get(url, headers=headers, timeout=30)
        if r.status_code >= 300:
            raise RuntimeError(f"Files API poll failed: HTTP {r.status_code}")
        state = r.json().get("state", "")
        if state == "ACTIVE":
            return
        if state == "FAILED":
            raise RuntimeError("Gemini could not process the uploaded file")
    raise RuntimeError("Gemini is still processing the file — try again in a moment")
